import re
import pdfplumber
from docx import Document

# ==========================================
# SHARED UTILITIES
# ==========================================

def clean_text(text):
    """Cleans up newlines and excessive whitespace from cell text."""
    if not text: return ""
    return re.sub(r'\s+', ' ', str(text)).strip()

# ==========================================
# DOCX PARSER (Unchanged - 100% Accurate)
# ==========================================

def process_docx(file_path):
    doc = Document(file_path)
    if not doc.tables: return []
    
    biggest_table = max(doc.tables, key=lambda t: len(t.rows) * len(t.columns))
    grid = [[clean_text(cell.text) for cell in row.cells] for row in biggest_table.rows]
    
    header_idx = -1
    for i, row in enumerate(grid):
        if any("MONDAY" in str(cell).upper() for cell in row):
            header_idx = i
            break
            
    if header_idx == -1: return []
    headers = [str(c).strip().upper() for c in grid[header_idx]]
    day_columns = {h: idx for idx, h in enumerate(headers) if h in ["MONDAY", "TUESDAY", "WEDNESDAY", "THURSDAY", "FRIDAY", "SATURDAY"]}

    completed_blocks = []
    active_blocks = {day: None for day in day_columns.keys()}

    for row_idx in range(header_idx + 1, len(grid)):
        row = grid[row_idx]
        if not row: continue
        
        time_col = str(row[0]).upper()
        if "SUMMARY" in time_col or "REGULAR TEACHING" in time_col: break

        matches = re.findall(r'(\d{1,2})[^\dA-Za-z]*(\d{2})\s*(AM|PM)', str(time_col), re.IGNORECASE)
        times = []
        for h, m, p in matches:
            hour = int(h)
            if hour == 12: hour = 12
            elif p.upper() == 'PM' and hour < 12: hour += 12
            times.append(f"{hour:02d}:{m}")
            
        start_time = times[0] if len(times) >= 1 else None
        end_time = times[-1] if len(times) >= 2 else start_time
        
        if not start_time or not end_time: continue

        for day, col_idx in day_columns.items():
            if col_idx >= len(row): continue
            cell_text = row[col_idx]
            is_empty = not cell_text or cell_text.lower() in ['break', 'lunch', 'n/a', '']

            if not is_empty:
                if active_blocks[day] and active_blocks[day]['event'] == cell_text:
                    active_blocks[day]['end'] = end_time
                else:
                    if active_blocks[day]: completed_blocks.append(active_blocks[day])
                    active_blocks[day] = {'day': day, 'event': cell_text, 'start': start_time, 'end': end_time}
            else:
                if active_blocks[day]:
                    completed_blocks.append(active_blocks[day])
                    active_blocks[day] = None

    for block in active_blocks.values():
        if block: completed_blocks.append(block)

    return completed_blocks

# ==========================================
# PDF PARSER (Color-Enforced Boundary Logic)
# ==========================================

def process_pdf(file_path):
    completed_blocks = []
    
    with pdfplumber.open(file_path) as pdf:
        page = pdf.pages[0]
        words = page.extract_words()
        
        headers = ["MONDAY", "TUESDAY", "WEDNESDAY", "THURSDAY", "FRIDAY", "SATURDAY"]
        header_words = [w for w in words if w['text'].upper() in headers]
        if not header_words: return []
        
        # 1. Map Columns using Midpoints
        col_cx = {w['text'].upper(): (w['x0'] + w['x1']) / 2 for w in header_words}
        sorted_days = sorted(col_cx.keys(), key=lambda d: col_cx[d])
        monday_x0 = next(w['x0'] for w in header_words if w['text'].upper() == "MONDAY")
        
        col_bounds = {}
        for i, day in enumerate(sorted_days):
            start_x = monday_x0 - 15 if i == 0 else (col_cx[sorted_days[i-1]] + col_cx[day]) / 2
            end_x = (col_cx[day] + col_cx[sorted_days[i+1]]) / 2 if i+1 < len(sorted_days) else page.width
            col_bounds[day] = (start_x, end_x)
            
        # 2. Strict Boundary for the Main Grid
        table_top = min(w['top'] for w in header_words) - 5
        summary_words = [w for w in words if "SUMMARY" in w['text'].upper() or "TEACHING" in w['text'].upper()]
        table_bottom = min(w['top'] for w in summary_words) - 5 if summary_words else page.height

        table_words = [w for w in words if w['top'] >= table_top and w['bottom'] <= table_bottom]

        # 3. Y-axis Uniform Time Grid
        time_slots = [
            "07:00", "07:30", "08:00", "08:30", "09:00", "09:30", "10:00", "10:30",
            "11:00", "11:30", "12:00", "12:30", "13:00", "13:30", "14:00", "14:30",
            "15:00", "15:30", "16:00", "16:30", "17:00", "17:30", "18:00", "18:30",
            "19:00", "19:30", "20:00"
        ]

        time_words = [w for w in table_words if w['x1'] < monday_x0 and re.search(r'\d{1,2}:\d{2}', w['text'])]
        time_words.sort(key=lambda w: w['top'])
        if not time_words: return []

        grid_top = time_words[0]['top']
        grid_bottom = time_words[-1]['bottom']
        slot_height = (grid_bottom - grid_top) / 26.0

        def snap_to_time(y):
            idx = int(round((y - grid_top) / slot_height))
            idx = max(0, min(26, idx))
            return time_slots[idx]

        # 4. Extract colored background blocks & Remove duplicates
        raw_rects = [r for r in page.rects if r['top'] >= grid_top - 5 and r['bottom'] <= grid_bottom + 5 and r['height'] > 5 and 20 < r['width'] < (page.width * 0.4)]
        
        unique_rects = []
        for r in raw_rects:
            r_center_x = (r['x0'] + r['x1']) / 2
            r_center_y = (r['top'] + r['bottom']) / 2
            is_dup = False
            for idx, ur in enumerate(unique_rects):
                ur_center_x = (ur['x0'] + ur['x1']) / 2
                ur_center_y = (ur['top'] + ur['bottom']) / 2
                if abs(r_center_x - ur_center_x) < 5 and abs(r_center_y - ur_center_y) < 5:
                    is_dup = True
                    if (r['width'] * r['height']) > (ur['width'] * ur['height']):
                        unique_rects[idx] = r
                    break
            if not is_dup:
                unique_rects.append(r)

        # 5. Assign every word a tag mapping it to its specific Colored Box
        def get_rect_idx(w):
            w_cx = (w['x0'] + w['x1']) / 2
            w_cy = (w['top'] + w['bottom']) / 2
            for idx, r in enumerate(unique_rects):
                if r['x0'] - 5 <= w_cx <= r['x1'] + 5 and r['top'] - 5 <= w_cy <= r['bottom'] + 5:
                    return idx
            return -1

        for w in table_words:
            w['rect_idx'] = get_rect_idx(w)

        # 6. Read data strictly within columns
        for day, (x0, x1) in col_bounds.items():
            col_words = []
            for w in table_words:
                w_cx = (w['x0'] + w['x1']) / 2
                if x0 <= w_cx < x1 and w['top'] > grid_top - 5:
                    text_upper = clean_text(w['text']).upper()
                    if text_upper not in headers and text_upper not in ['BREAK', 'LUNCH', 'N/A', '']:
                        col_words.append(w)
                        
            col_words.sort(key=lambda w: w['top'])
            
            # Cluster floating text strictly restricted by colored boxes
            clusters = []
            current_cluster = []
            for w in col_words:
                if not current_cluster:
                    current_cluster.append(w)
                else:
                    last_w = max(current_cluster, key=lambda x: x['bottom'])
                    
                    # STRICT RULE: They MUST share the exact same colored box to merge!
                    same_rect = (w['rect_idx'] == last_w['rect_idx'])
                    
                    # If they are in the same colored box, give generous merge room (25px)
                    if same_rect and w['rect_idx'] != -1 and w['top'] - last_w['bottom'] <= 25:
                        current_cluster.append(w)
                    # If they both have NO background box, use tight spacing merge (12px)
                    elif w['rect_idx'] == -1 and last_w['rect_idx'] == -1 and w['top'] - last_w['bottom'] <= 12:
                        current_cluster.append(w)
                    # Force Split (Solves the "Community / Consultation" merging bug)
                    else:
                        clusters.append(current_cluster)
                        current_cluster = [w]
                        
            if current_cluster: clusters.append(current_cluster)

            parsed_clusters = []
            for c in clusters:
                c.sort(key=lambda w: (w['top'], w['x0']))
                text = clean_text(" ".join([w['text'] for w in c]))
                if not text: continue
                
                parsed_clusters.append({
                    'text': text,
                    'top': min(w['top'] for w in c),
                    'bottom': max(w['bottom'] for w in c),
                    'center_y': (min(w['top'] for w in c) + max(w['bottom'] for w in c)) / 2,
                    'rect_idx': c[0]['rect_idx'] # Pass the ID of the colored box forward
                })

            for i, cluster in enumerate(parsed_clusters):
                matched_rect = None
                # Retrieve the boundaries of the specific colored box to set the exact times
                if cluster['rect_idx'] != -1:
                    matched_rect = unique_rects[cluster['rect_idx']]

                if matched_rect:
                    start_time = snap_to_time(matched_rect['top'])
                    end_time = snap_to_time(matched_rect['bottom'])
                else:
                    # Fallback (If white background): Find midpoint between classes and snap to grid
                    start_y = grid_top if i == 0 else (parsed_clusters[i-1]['center_y'] + cluster['center_y']) / 2
                    end_y = grid_bottom if i == len(parsed_clusters) - 1 else (cluster['center_y'] + parsed_clusters[i+1]['center_y']) / 2
                    start_time = snap_to_time(start_y)
                    end_time = snap_to_time(end_y)
                
                if start_time != end_time:
                    completed_blocks.append({
                        'day': day,
                        'event': cluster['text'],
                        'start': start_time,
                        'end': end_time
                    })
                    
    return completed_blocks

# ==========================================
# MAIN EXECUTION
# ==========================================

def run():
    files = ["Engr. Carlo.docx", "Engr. Carlo.pdf", "new_sched_1.pdf"]
    day_order = {"MONDAY": 1, "TUESDAY": 2, "WEDNESDAY": 3, "THURSDAY": 4, "FRIDAY": 5, "SATURDAY": 6}

    for file_path in files:
        try:
            print(f"\nExtracted from {file_path}:")
            print("=" * 50)
            
            if file_path.endswith('.docx'):
                schedules = process_docx(file_path)
            else:
                schedules = process_pdf(file_path)
                
            if not schedules:
                print("No valid schedule data extracted.")
                continue

            schedules.sort(key=lambda x: (day_order.get(x['day'], 99), x['start']))
            
            for block in schedules:
                print(f"DAY:   {block['day']}")
                print(f"TIME:  {block['start']} to {block['end']}")
                print(f"EVENT: {block['event']}")
                print("-" * 50)
                
        except Exception as e:
            print(f"Failed to process {file_path}: {e}")

if __name__ == "__main__":
    run()